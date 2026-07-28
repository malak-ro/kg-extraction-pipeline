"""
Chargement de documents (PDF / TXT / DOCX) derrière une interface commune.

Pourquoi une interface plutôt que 3 fonctions séparées (load_txt, load_pdf, load_docx)
appelées avec des if/elif sur l'extension ?

- Open/Closed Principle : ajouter un futur format (HTML, Markdown...) = ajouter
  UNE classe, sans toucher au code existant. Avec des if/elif, chaque nouveau
  format oblige à modifier une fonction déjà en place — plus risqué.
- Le code appelant (le pipeline d'ingestion) ne connaît que `DocumentLoader` /
  `DocumentLoaderRegistry` — jamais `PdfLoader` ou `pdfplumber` directement
  (Dependency Inversion). On pourra remplacer pdfplumber par un autre moteur
  PDF plus tard sans casser quoi que ce soit en amont.
"""
from __future__ import annotations

from abc import ABC, abstractmethod
from dataclasses import dataclass
from pathlib import Path

import docx
import pdfplumber

from src.utils.exceptions import DocumentLoadError
from src.utils.logger import get_logger

logger = get_logger(__name__)


@dataclass
class LoadedDocument:
    """Représentation uniforme d'un document, quel que soit son format d'origine.
    Tout le reste du pipeline (nettoyage, NER...) travaille sur CET objet,
    jamais sur un PDF ou un DOCX directement."""

    source_path: Path
    text: str
    doc_format: str  # "txt" | "pdf" | "docx"

    @property
    def char_count(self) -> int:
        return len(self.text)


class DocumentLoader(ABC):
    """Contrat commun à tous les loaders."""

    @abstractmethod
    def supports(self, path: Path) -> bool:
        """True si CE loader sait lire ce fichier (généralement : l'extension)."""

    @abstractmethod
    def load(self, path: Path) -> LoadedDocument:
        """Lit le fichier et retourne un LoadedDocument.
        Lève DocumentLoadError si la lecture échoue, pour que l'appelant
        n'ait qu'un seul type d'exception à gérer, peu importe le format."""


class TxtLoader(DocumentLoader):
    def supports(self, path: Path) -> bool:
        return path.suffix.lower() == ".txt"

    def load(self, path: Path) -> LoadedDocument:
        try:
            text = path.read_text(encoding="utf-8")
        except (OSError, UnicodeDecodeError) as exc:
            raise DocumentLoadError(f"Impossible de lire {path} : {exc}") from exc
        return LoadedDocument(source_path=path, text=text, doc_format="txt")


class PdfLoader(DocumentLoader):
    def supports(self, path: Path) -> bool:
        return path.suffix.lower() == ".pdf"

    def load(self, path: Path) -> LoadedDocument:
        try:
            with pdfplumber.open(path) as pdf:
                # extract_text() renvoie None sur une page sans texte détectable
                # (page scannée/image) — le `or ""` évite un crash sur ce cas.
                pages = [page.extract_text() or "" for page in pdf.pages]
        except Exception as exc:  # pdfplumber peut lever plusieurs types selon le PDF
            raise DocumentLoadError(f"Impossible de lire {path} : {exc}") from exc

        text = "\n".join(pages)
        if not text.strip():
            logger.warning(
                "Aucun texte extrait de %s — probablement un PDF scanné (image) "
                "sans OCR, pas un vrai échec de lecture.",
                path.name,
            )
        return LoadedDocument(source_path=path, text=text, doc_format="pdf")


class DocxLoader(DocumentLoader):
    def supports(self, path: Path) -> bool:
        return path.suffix.lower() == ".docx"

    def load(self, path: Path) -> LoadedDocument:
        try:
            document = docx.Document(path)
            text = "\n".join(p.text for p in document.paragraphs)
        except Exception as exc:
            raise DocumentLoadError(f"Impossible de lire {path} : {exc}") from exc
        return LoadedDocument(source_path=path, text=text, doc_format="docx")


class DocumentLoaderRegistry:
    """Essaie chaque loader enregistré jusqu'à trouver celui qui convient.
    Ajouter un format = ajouter une ligne dans `_loaders`, rien d'autre à changer
    ailleurs dans le projet (c'est concrètement l'Open/Closed Principle)."""

    def __init__(self) -> None:
        self._loaders: list[DocumentLoader] = [TxtLoader(), PdfLoader(), DocxLoader()]

    def load(self, path: Path) -> LoadedDocument:
        for loader in self._loaders:
            if loader.supports(path):
                logger.info("Chargement de %s via %s", path.name, type(loader).__name__)
                return loader.load(path)
        raise DocumentLoadError(f"Aucun loader disponible pour le format de {path}")
